import fitz
import docx
from io import BytesIO

def extract_text_from_pdf(file):
    """Extract text from a PDF file-like object."""
    text = ""
    try:
        # fitz.open can accept bytes or a file-like object
        if hasattr(file, 'read'):
            file_bytes = file.read()
            pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
        else:
            pdf_document = fitz.open(file)
        for page in pdf_document:
            text += page.get_text()
        pdf_document.close()
    except Exception as e:
        print(f"Error reading PDF file: {e}")
    # Normalize text by replacing multiple newlines with a single space
    import re
    text = re.sub(r'\n{2,}', '\n\n', text)
    text = re.sub(r'(?<!\n)\n(?!\n)', ' ', text)
    text = text.strip()
    return text

def extract_text_from_docx(file):
    """Extract text from a DOCX file-like object."""
    text = ""
    try:
        if hasattr(file, 'read'):
            file_bytes = file.read()
            file_stream = BytesIO(file_bytes)
            doc = docx.Document(file_stream)
        else:
            doc = docx.Document(file)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
    except Exception as e:
        print(f"Error reading DOCX file: {e}")
    # Normalize text by replacing multiple newlines with a single space
    import re
    text = re.sub(r'\n{2,}', '\n\n', text)
    text = re.sub(r'(?<!\n)\n(?!\n)', ' ', text)
    text = text.strip()
    return text

def extract_text_from_file(file, filename):
    """Extract text from a file-like object based on its extension."""
    if filename.lower().endswith('.pdf'):
        return extract_text_from_pdf(file)
    elif filename.lower().endswith('.docx'):
        return extract_text_from_docx(file)
    else:
        raise ValueError("Unsupported file type. Please provide a PDF or DOCX file.")